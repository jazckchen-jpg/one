"""
文档生成工具 —— 支持 Markdown、Word 格式报告
"""

import os
from datetime import datetime
from typing import Optional


class DocumentGenerator:
    """文档生成器，支持多种输出格式"""

    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def save_markdown(self, content: str, filename: str) -> str:
        """
        保存为 Markdown 文件

        Args:
            content: Markdown 内容
            filename: 文件名（不含扩展名）

        Returns:
            文件路径
        """
        if not filename.endswith(".md"):
            filename += ".md"
        filepath = os.path.join(self.output_dir, filename)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"  [OK] Markdown 报告已保存: {filepath}")
        return filepath

    def save_word(self, markdown_content: str, filename: str) -> str:
        """
        将 Markdown 转换为 Word 文档并保存

        Args:
            markdown_content: Markdown 格式内容
            filename: 文件名（不含扩展名）

        Returns:
            文件路径
        """
        if not filename.endswith(".docx"):
            filename += ".docx"
        filepath = os.path.join(self.output_dir, filename)

        try:
            # 延迟导入，避免没有 python-docx 时出错
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 解析 Markdown 行并逐行转换
            lines = markdown_content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]

                # 标题 (## 或 #)
                if line.startswith("# ") or line.startswith("## "):
                    level = 1 if line.startswith("# ") else 2
                    text = line.lstrip("#").strip()
                    if level == 1:
                        heading = doc.add_heading(text, level=1)
                    else:
                        heading = doc.add_heading(text, level=2)
                    # 标题前空行
                    i += 1
                    continue

                # 分割线
                if line.strip() == "---":
                    doc.add_paragraph("_" * 60)
                    i += 1
                    continue

                # 列表项
                if line.startswith("- ") or line.startswith("* "):
                    text = line[2:].strip()
                    doc.add_paragraph(text, style="List Bullet")
                    i += 1
                    continue

                # 数字列表
                if line and line[0].isdigit() and ". " in line[:4]:
                    text = line.split(". ", 1)[1] if ". " in line else line
                    doc.add_paragraph(text, style="List Number")
                    i += 1
                    continue

                # 空行
                if not line.strip():
                    i += 1
                    continue

                # 普通段落
                # 处理加粗
                text = line.strip()
                if text:
                    p = doc.add_paragraph(text)

                i += 1

            # 设置默认字体
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(12)

            doc.save(filepath)
            print(f"  [OK] Word 文档已保存: {filepath}")

        except ImportError:
            # 没有 python-docx 时降级为 Markdown
            print("  [!] 未安装 python-docx，降级保存为 Markdown 格式")
            filepath = self.save_markdown(markdown_content, filename.replace(".docx", ".md"))

        return filepath

    def build_report(self, title: str, sections: list[dict], format: str = "markdown") -> str:
        """
        构建完整报告

        Args:
            title: 报告标题
            sections: 章节列表，每项包含 title, content, level
            format: 输出格式 (markdown/docx)

        Returns:
            输出文件路径
        """
        # 生成 Markdown 内容
        md_lines = [
            f"# {title}",
            "",
            f"> 生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
            "",
            "---",
            "",
            "## 目录",
            "",
        ]

        # 生成目录
        for i, section in enumerate(sections, 1):
            md_lines.append(f"{i}. [{section.get('title', '')}](#section-{i})")

        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

        # 生成正文
        for i, section in enumerate(sections):
            title = section.get("title", f"第{i+1}章")
            content = section.get("content", "")
            level = section.get("level", 2)

            md_lines.append(f"{'#' * level} {title}")
            md_lines.append("")
            md_lines.append(content)
            md_lines.append("")

        # 添加页脚
        md_lines.append("---")
        md_lines.append("")
        md_lines.append(
            f"*本报告由「资料搜集与报告整理智能体」自动生成*  \n"
            f"*生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*"
        )

        md_content = "\n".join(md_lines)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = title.replace(" ", "_").replace("/", "_")[:30]

        if format == "docx":
            return self.save_word(md_content, f"{safe_title}_{timestamp}")
        else:
            return self.save_markdown(md_content, f"{safe_title}_{timestamp}")
